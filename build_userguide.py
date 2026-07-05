"""Generate a Word (.docx) user guide for Chronicle (the conversation viewer).

Run: python build_userguide.py
Produces: "Chronicle - User Guide.docx"
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor

ACCENT = RGBColor(0xB8, 0x55, 0x2A)   # terracotta, matches the viewer
INK = RGBColor(0x2B, 0x2A, 0x27)
MUTED = RGBColor(0x6B, 0x66, 0x60)
OUT = Path(__file__).parent / "Chronicle - User Guide.docx"


def main() -> None:
    doc = Document()

    # Base font
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK

    # ---- Title ----
    title = doc.add_paragraph()
    run = title.add_run("Chronicle: Browsing Your Chat History")
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.color.rgb = ACCENT
    sub = doc.add_paragraph()
    r = sub.add_run("A simple guide to opening and searching your exported Claude or ChatGPT data")
    r.font.size = Pt(13)
    r.font.color.rgb = MUTED
    doc.add_paragraph()

    def h(text: str) -> None:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(15)
        run.font.bold = True
        run.font.color.rgb = ACCENT
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(4)

    def body(text: str) -> None:
        p = doc.add_paragraph(text)
        p.paragraph_format.space_after = Pt(6)

    def step(num: int, lead: str, rest: str) -> None:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        n = p.add_run(f"{num}.  ")
        n.font.bold = True
        n.font.color.rgb = ACCENT
        b = p.add_run(lead + " ")
        b.font.bold = True
        p.add_run(rest)

    def bullet(lead: str, rest: str = "") -> None:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        if lead:
            b = p.add_run(lead + (" " if rest else ""))
            b.font.bold = True
        if rest:
            p.add_run(rest)

    def note(label: str, text: str) -> None:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        lab = p.add_run(label + "  ")
        lab.font.bold = True
        lab.font.color.rgb = ACCENT
        t = p.add_run(text)
        t.font.italic = True
        t.font.color.rgb = MUTED

    # ---- Intro ----
    body(
        "Chronicle is a viewer file named conversations-viewer.html that lets you read "
        "and search your past Claude or ChatGPT conversations in a familiar, chat-style "
        "layout. You load it with the .zip data export from either service. It runs "
        "entirely in your web browser."
    )
    note(
        "Your data stays private.",
        "Everything happens on your own computer. Nothing is uploaded, there is no "
        "server, and it works without an internet connection. Your data never leaves "
        "your device.",
    )

    # ---- How to download your data ----
    h("How to download your data")
    body("First, export your history from the service you use:")
    bullet("Claude:", "go to claude.ai, then Settings → Privacy → Export data. You "
           "will receive a .zip by email.")
    bullet("ChatGPT:", "go to chatgpt.com, then Settings → Data controls → Export "
           "data. You will receive a .zip by email.")
    body("Either .zip works with Chronicle, and you do not need to unzip it.")

    # ---- Getting started ----
    h("Getting started")
    step(1, "Save both files together.",
         "Put conversations-viewer.html and your .zip file in the same folder (for "
         "example, your Desktop). You do not need to unzip anything.")
    step(2, "Open the viewer.",
         "Double-click conversations-viewer.html. It opens in your default web "
         "browser (Chrome, Safari, Edge, or Firefox all work).")
    step(3, "Load your data.",
         "Drag your .zip file anywhere onto the page, or click the “Choose export…” "
         "button and select the .zip. Chronicle recognizes both Claude and ChatGPT "
         "exports automatically — your conversations (and, for Claude, any projects) "
         "load on their own.")
    step(4, "Start browsing.",
         "Your conversations appear in the list on the left, newest first. Click any "
         "one to read it.")
    note(
        "Large files take a moment.",
        "If your history is big, give it a few seconds to load after you select the "
        "file — it is reading everything on your computer.",
    )

    # ---- Finding things ----
    h("Finding a conversation")
    bullet("Search:", "Type in the search box at the top left to search across every "
           "message. Matches are highlighted, and each result shows how many times "
           "your term appears.")
    bullet("Sort:", "Use the “sort” dropdown to order the list by most recent, oldest "
           "first, or title.")
    bullet("Open:", "Click a conversation to read the full back-and-forth. Your "
           "messages appear on the right, the assistant’s (Claude or ChatGPT) on the "
           "left.")

    # ---- Artifacts, images, and generated files ----
    h("Things the assistant made for you")
    body(
        "Content the assistant produced appears inside the conversation, either as its "
        "own card or inline:"
    )
    bullet("Claude artifacts and files", "(documents, web pages, diagrams) render so "
           "you can read them directly, each with a “Copy” button.")
    bullet("ChatGPT images", "(pictures you sent, or ones that were generated for you) "
           "appear inline in the conversation.")
    body(
        "Some Claude cards show computer code rather than a finished document. This is "
        "normal and expected — see the next section."
    )

    # ---- Browsing by artifact ----
    h("The Artifacts tab")
    body(
        "For Claude exports, the “Artifacts” tab gathers everything Claude built — "
        "artifacts and generated files — into one searchable place, so you can browse "
        "them without opening each conversation. Each one links back to the chat it "
        "came from, and has its own “Export” button to save it as a file. (ChatGPT "
        "exports have no artifacts, so this tab stays empty for them.)"
    )

    # ---- Projects ----
    h("Projects (Claude only)")
    body(
        "If you used Projects in Claude, switch to the “Projects” tab at the top of "
        "the left panel. Your projects load from the same .zip file. Each project "
        "shows its custom instructions and its documents. Click a document to expand "
        "it and read the text. ChatGPT exports don’t include projects, so this tab "
        "stays empty for them."
    )

    # ---- Saving a conversation ----
    h("Saving or sharing a conversation")
    body(
        "At the top of any open conversation there are two buttons:"
    )
    bullet("Copy", "puts the whole conversation on your clipboard as plain text.")
    bullet("Export .md", "saves the conversation as a Markdown file, named with its "
           "date and title.")

    # ---- What you'll see / important to understand ----
    h("Why some things look like raw code")
    body(
        "When the assistant produced a Word, Excel, PowerPoint, or PDF file, your "
        "export stores the code it wrote to generate that file — not the finished file "
        "itself. So you will see code instead of the polished document."
    )
    body("To get the finished document back:")
    step(1, "Click the “Copy” button", "on the card.")
    step(2, "Open a new chat with the same assistant", "(claude.ai or chatgpt.com), "
         "paste what you copied, and send the message “Run this.”")
    step(3, "The assistant regenerates the file", "for you to download.")

    # ---- What's included ----
    h("What is and isn’t in your export")
    bullet("Included:", "all your conversations and the assistant’s answers; for "
           "Claude, project instructions and the extracted text of documents; for "
           "ChatGPT, the images from your chats.")
    bullet("Not included:", "the original copies of files you uploaded, and the "
           "finished documents the assistant generated (Word, Excel, PDF, slides). The "
           "export keeps the text and the code, but not those original or final files.")

    # ---- Tips ----
    h("Tips and troubleshooting")
    bullet("Nothing happens when I drag the file:", "use the “Choose export…” button "
           "instead — it always works.")
    bullet("I want to look at a different export:", "click “↻ load another” at the top "
           "left and pick a different .zip.")
    bullet("The info banner:", "the note at the top of a conversation can be dismissed "
           "with “Got it” and won’t come back.")
    bullet("It looks empty:", "make sure you selected the .zip file itself, not an "
           "unzipped folder.")
    bullet("My ChatGPT export has several files:", "that’s normal — ChatGPT splits "
           "conversations across multiple files inside the .zip. Just load the whole "
           ".zip and Chronicle combines them.")

    doc.add_paragraph()
    closing = doc.add_paragraph()
    c = closing.add_run("That’s it — open the viewer, load your zip, and explore.")
    c.font.italic = True
    c.font.color.rgb = MUTED

    doc.save(OUT)
    print(f"Wrote {OUT.name}")


if __name__ == "__main__":
    main()
